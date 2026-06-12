import os
import re
import time
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import ImageGrab


TEMPLATE_FILE = "Template.docx"
OUTPUT_DIR = "output"
TEMP_IMG_DIR = "temp_clipboard_images"


# -----------------------------
# General helpers
# -----------------------------
def normalize_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()


def safe_filename(s: str) -> str:
    s = re.sub(r'[\\/*?:"<>|]+', "_", s)
    return s.strip().replace(" ", "_")


def ensure_dirs():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(TEMP_IMG_DIR, exist_ok=True)


def iter_all_paragraphs(container):
    for p in container.paragraphs:
        yield p
    for t in container.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from iter_all_paragraphs(cell)


def paragraph_full_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def set_paragraph_text(paragraph, text: str):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


# -----------------------------
# Replacement helpers
# -----------------------------
def replace_text_in_paragraph(paragraph, replacements):
    full_text = paragraph_full_text(paragraph)
    new_text = full_text
    for old, new in replacements.items():
        if old:
            new_text = new_text.replace(old, new)
    if new_text != full_text:
        set_paragraph_text(paragraph, new_text)


def replace_text_in_container(container, replacements):
    for p in iter_all_paragraphs(container):
        replace_text_in_paragraph(p, replacements)


def replace_everywhere(doc, replacements):
    replace_text_in_container(doc, replacements)
    for section in doc.sections:
        replace_text_in_container(section.header, replacements)
        replace_text_in_container(section.footer, replacements)


# -----------------------------
# Anchor search
# -----------------------------
def find_paragraph_by_contains(container, needles, case_insensitive=True):
    for p in iter_all_paragraphs(container):
        txt = paragraph_full_text(p)
        check = txt.lower() if case_insensitive else txt
        for needle in needles:
            nd = needle.lower() if case_insensitive else needle
            if nd in check:
                return p
    return None


def find_any_anchor(doc, needles):
    p = find_paragraph_by_contains(doc, needles)
    if p:
        return p, "body"

    for i, section in enumerate(doc.sections):
        p = find_paragraph_by_contains(section.header, needles)
        if p:
            return p, f"header_{i}"
        p = find_paragraph_by_contains(section.footer, needles)
        if p:
            return p, f"footer_{i}"

    return None, None


# -----------------------------
# Insert paragraph/image
# -----------------------------
def insert_paragraph_after(paragraph, text="", style=None):
    new_para = paragraph._parent.add_paragraph()
    paragraph._p.addnext(new_para._p)
    if text:
        new_para.add_run(text)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def insert_image_after(paragraph, image_path, width=Inches(4.5), caption=None):
    p_img = insert_paragraph_after(paragraph)
    run = p_img.add_run()
    run.add_picture(image_path, width=width)

    if caption:
        p_cap = insert_paragraph_after(p_img, caption)
        try:
            p_cap.runs[0].italic = True
        except Exception:
            pass
        return p_img, p_cap
    return p_img, None


# -----------------------------
# Hyperlink helper
# -----------------------------
def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


# -----------------------------
# Clipboard image support
# -----------------------------
def save_clipboard_image(slot_name):
    """
    Reads image from clipboard and saves it to temp folder.
    Works best on Windows.
    """
    ensure_dirs()
    img = ImageGrab.grabclipboard()

    if img is None:
        raise ValueError("No image found in clipboard. Copy an image first, then click Paste.")

    # Sometimes clipboard contains file paths instead of image object
    if isinstance(img, list):
        raise ValueError("Clipboard contains file list, not an image. Please copy an actual image.")

    filename = f"{slot_name}_{int(time.time())}.png"
    path = os.path.join(TEMP_IMG_DIR, filename)
    img.save(path, "PNG")
    return path


# -----------------------------
# Business logic
# -----------------------------
def build_olt_label(equipment: str, custom_olt_label: str) -> str:
    eq = normalize_spaces(equipment).lower()
    if eq == "nokia lightspan mf-2":
        return "OLT MF-2"
    return normalize_spaces(custom_olt_label) or normalize_spaces(equipment)


def build_fuse_line(olt_label: str, equipment: str) -> str:
    return f"FUSE No: L3 {olt_label} – ({normalize_spaces(equipment)} Power tapping point)"


def insert_power_section_content(doc, data):
    fuse_variants = [
        "FUSE No: L3 Nokia OLT MF-2 – ( Nokia Power tapping point)",
        "FUSE No: L3 Nokia OLT MF-2 – (Nokia Power tapping point)",
        "FUSE No: L3 Nokia OLT MF-2",
        "FUSE No:"
    ]

    anchor, where = find_any_anchor(doc, fuse_variants)
    if not anchor:
        return False, "Could not find FUSE anchor text in template."

    olt_label = build_olt_label(data["equipment"], data["olt_label_custom"])
    new_fuse = build_fuse_line(olt_label, data["equipment"])

    old_text = paragraph_full_text(anchor)
    replaced = old_text
    exacts = [
        "FUSE No: L3 Nokia OLT MF-2 – ( Nokia Power tapping point)",
        "FUSE No: L3 Nokia OLT MF-2 – (Nokia Power tapping point)",
    ]
    exact_hit = False
    for ex in exacts:
        if ex in replaced:
            replaced = replaced.replace(ex, new_fuse)
            exact_hit = True

    if exact_hit:
        set_paragraph_text(anchor, replaced)
    else:
        set_paragraph_text(anchor, new_fuse)

    rs1_line = f"RS1 + {data['rs1_rectifier_name']} + {data['rs1_load_assignment']}"
    rs2_line = f"RS2 + {data['rs2_rectifier_name']} + {data['rs2_load_assignment']}"

    p1 = insert_paragraph_after(anchor, rs1_line)
    last = p1
    if data.get("rs1_image"):
        img_p, _ = insert_image_after(p1, data["rs1_image"], width=Inches(4.5), caption="RS1")
        last = img_p

    p2 = insert_paragraph_after(last, rs2_line)
    last = p2
    if data.get("rs2_image"):
        img_p, _ = insert_image_after(p2, data["rs2_image"], width=Inches(4.5), caption="RS2")
        last = img_p

    return True, f"Inserted RS1/RS2 content under FUSE anchor in {where}."


def insert_supporting_documents(doc, pdf_path):
    if not pdf_path:
        return False, "No TSSR PDF selected."

    anchor, where = find_any_anchor(doc, ["Supporting Documents"])
    if not anchor:
        return False, "Could not find 'Supporting Documents' section."

    p = insert_paragraph_after(anchor)
    p.add_run("TSSR: ")
    file_url = "file:///" + pdf_path.replace("\\", "/")
    add_hyperlink(p, os.path.basename(pdf_path), file_url)
    return True, f"Inserted TSSR hyperlink under Supporting Documents in {where}."


def insert_existing_rectifier_image(doc, image_path):
    if not image_path:
        return False, "No Page 8 rectifier image selected."

    anchor, where = find_any_anchor(doc, ["Existing Rectifier"])
    if not anchor:
        return False, "Could not find 'Existing Rectifier' section."

    insert_image_after(anchor, image_path, width=Inches(5.0), caption="Existing Rectifier")
    return True, f"Inserted rectifier image under Existing Rectifier in {where}."


def generate_mop(data):
    if not os.path.exists(TEMPLATE_FILE):
        raise FileNotFoundError(f"Template file not found: {TEMPLATE_FILE}")

    ensure_dirs()
    doc = Document(TEMPLATE_FILE)

    replacements = {
        "CDO-604": data["site_name"],
        "MIN699": data["plaid"],
        "Nokia Lightspan MF-2": data["equipment"],
        "John Carlo Rabanes": data["prepared_by"],
        "OLT Rollout Engineer": data["position"],
        "< May 19- June 19, 2026 10:00AM-6:00PM>": data["target_datetime"],
        "May 19- June 19, 2026 10:00AM-6:00PM": data["target_datetime"],
    }

    replace_everywhere(doc, replacements)

    power_ok, power_msg = insert_power_section_content(doc, data)
    pdf_ok, pdf_msg = insert_supporting_documents(doc, data.get("tssr_pdf"))
    rect_ok, rect_msg = insert_existing_rectifier_image(doc, data.get("rectifier_image_page8"))

    output_file = os.path.join(
        OUTPUT_DIR,
        f"MOP_{safe_filename(data['site_name'])}_{safe_filename(data['plaid'])}.docx"
    )
    doc.save(output_file)

    return output_file, [power_msg, pdf_msg, rect_msg]


# -----------------------------
# GUI
# -----------------------------
class MopApp:
    def __init__(self, root):
        self.root = root
        self.root.title("MOP Automation with Clipboard Paste")
        self.root.geometry("980x700")

        self.entries = {}
        self.file_paths = {
            "rs1_image": "",
            "rs2_image": "",
            "rectifier_image_page8": "",
            "tssr_pdf": "",
        }

        main = tk.Frame(root)
        main.pack(fill="both", expand=True, padx=10, pady=10)

        fields = [
            ("Site Name", "site_name"),
            ("Plaid", "plaid"),
            ("Equipment", "equipment"),
            ("Custom OLT Label (if not Nokia Lightspan MF-2)", "olt_label_custom"),
            ("Prepared By", "prepared_by"),
            ("Position", "position"),
            ("Target Date and Time of Implementation", "target_datetime"),
            ("RS1 Rectifier Name", "rs1_rectifier_name"),
            ("RS1 Load Assignment", "rs1_load_assignment"),
            ("RS2 Rectifier Name", "rs2_rectifier_name"),
            ("RS2 Load Assignment", "rs2_load_assignment"),
        ]

        row = 0
        for label_text, key in fields:
            tk.Label(main, text=label_text, anchor="w").grid(row=row, column=0, sticky="w", padx=5, pady=5)
            ent = tk.Entry(main, width=75)
            ent.grid(row=row, column=1, columnspan=3, sticky="we", padx=5, pady=5)
            self.entries[key] = ent
            row += 1

        self.entries["equipment"].insert(0, "Nokia Lightspan MF-2")
        self.entries["position"].insert(0, "OLT Rollout Engineer")
        self.entries["target_datetime"].insert(0, "May 19- June 19, 2026 10:00AM-6:00PM")

        # RS1
        tk.Label(main, text="RS1 Image").grid(row=row, column=0, sticky="w", padx=5, pady=5)
        tk.Button(main, text="Select File", command=lambda: self.select_file("rs1_image", [("Image Files", "*.png;*.jpg;*.jpeg;*.bmp")], self.rs1_label)).grid(row=row, column=1, sticky="w", padx=5)
        tk.Button(main, text="Paste Clipboard", command=lambda: self.paste_image("rs1_image", self.rs1_label)).grid(row=row, column=2, sticky="w", padx=5)
        self.rs1_label = tk.Label(main, text="No image selected", anchor="w")
        self.rs1_label.grid(row=row, column=3, sticky="w", padx=5)
        row += 1

        # RS2
        tk.Label(main, text="RS2 Image").grid(row=row, column=0, sticky="w", padx=5, pady=5)
        tk.Button(main, text="Select File", command=lambda: self.select_file("rs2_image", [("Image Files", "*.png;*.jpg;*.jpeg;*.bmp")], self.rs2_label)).grid(row=row, column=1, sticky="w", padx=5)
        tk.Button(main, text="Paste Clipboard", command=lambda: self.paste_image("rs2_image", self.rs2_label)).grid(row=row, column=2, sticky="w", padx=5)
        self.rs2_label = tk.Label(main, text="No image selected", anchor="w")
        self.rs2_label.grid(row=row, column=3, sticky="w", padx=5)
        row += 1

        # Rectifier
        tk.Label(main, text="Page 8 Existing Rectifier Image").grid(row=row, column=0, sticky="w", padx=5, pady=5)
        tk.Button(main, text="Select File", command=lambda: self.select_file("rectifier_image_page8", [("Image Files", "*.png;*.jpg;*.jpeg;*.bmp")], self.rectifier_label)).grid(row=row, column=1, sticky="w", padx=5)
        tk.Button(main, text="Paste Clipboard", command=lambda: self.paste_image("rectifier_image_page8", self.rectifier_label)).grid(row=row, column=2, sticky="w", padx=5)
        self.rectifier_label = tk.Label(main, text="No image selected", anchor="w")
        self.rectifier_label.grid(row=row, column=3, sticky="w", padx=5)
        row += 1

        # PDF
        tk.Label(main, text="TSSR PDF").grid(row=row, column=0, sticky="w", padx=5, pady=5)
        tk.Button(main, text="Select PDF", command=lambda: self.select_file("tssr_pdf", [("PDF Files", "*.pdf")], self.pdf_label)).grid(row=row, column=1, sticky="w", padx=5)
        self.pdf_label = tk.Label(main, text="No PDF selected", anchor="w")
        self.pdf_label.grid(row=row, column=3, sticky="w", padx=5)
        row += 1

        # Buttons
        btn_frame = tk.Frame(main)
        btn_frame.grid(row=row, column=0, columnspan=4, pady=20)

        tk.Button(btn_frame, text="Generate MOP", width=18, command=self.generate).pack(side="left", padx=8)
        tk.Button(btn_frame, text="Clear", width=12, command=self.clear_form).pack(side="left", padx=8)
        tk.Button(btn_frame, text="Exit", width=12, command=self.root.quit).pack(side="left", padx=8)

        # Instructions
        instructions = (
            "Clipboard paste usage:\n"
            "1. Copy an image to clipboard (Snipping Tool / screenshot / copied image).\n"
            "2. Click 'Paste Clipboard' for RS1, RS2, or Rectifier.\n"
            "3. The image is saved temporarily and used in the generated Word file.\n\n"
            "Notes:\n"
            "- No placeholders required.\n"
            "- Script searches for anchor text like 'FUSE No:', 'Supporting Documents', and 'Existing Rectifier'.\n"
            "- PDF is inserted as a clickable hyperlink."
        )
        tk.Label(main, text=instructions, justify="left", fg="gray25").grid(row=row+1, column=0, columnspan=4, sticky="w", padx=5, pady=10)

        main.columnconfigure(3, weight=1)

    def select_file(self, key, filetypes, label_widget):
        path = filedialog.askopenfilename(filetypes=filetypes)
        if path:
            self.file_paths[key] = path
            label_widget.config(text=os.path.basename(path))

    def paste_image(self, key, label_widget):
        try:
            path = save_clipboard_image(key)
            self.file_paths[key] = path
            label_widget.config(text=f"[Pasted] {os.path.basename(path)}")
            messagebox.showinfo("Clipboard Image", f"Image pasted successfully:\n{path}")
        except Exception as e:
            messagebox.showerror("Paste Failed", str(e))

    def clear_form(self):
        for entry in self.entries.values():
            entry.delete(0, tk.END)

        self.file_paths = {
            "rs1_image": "",
            "rs2_image": "",
            "rectifier_image_page8": "",
            "tssr_pdf": "",
        }

        self.rs1_label.config(text="No image selected")
        self.rs2_label.config(text="No image selected")
        self.rectifier_label.config(text="No image selected")
        self.pdf_label.config(text="No PDF selected")

        self.entries["equipment"].insert(0, "Nokia Lightspan MF-2")
        self.entries["position"].insert(0, "OLT Rollout Engineer")
        self.entries["target_datetime"].insert(0, "May 19- June 19, 2026 10:00AM-6:00PM")

    def generate(self):
        data = {k: e.get().strip() for k, e in self.entries.items()}
        data.update(self.file_paths)

        required = [
            "site_name",
            "plaid",
            "equipment",
            "prepared_by",
            "position",
            "target_datetime",
            "rs1_rectifier_name",
            "rs1_load_assignment",
            "rs2_rectifier_name",
            "rs2_load_assignment",
        ]
        missing = [field for field in required if not data.get(field)]
        if missing:
            messagebox.showerror("Missing Fields", "Please fill in:\n- " + "\n- ".join(missing))
            return

        try:
            output_file, notes = generate_mop(data)
            msg = "MOP generated successfully:\n\n" + output_file + "\n\nNotes:\n- " + "\n- ".join(notes)
            messagebox.showinfo("Success", msg)
        except Exception as e:
            messagebox.showerror("Error", str(e))


if __name__ == "__main__":
    ensure_dirs()
    root = tk.Tk()
    app = MopApp(root)
    root.mainloop()